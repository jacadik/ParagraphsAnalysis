# utils/document_processor.py
import os
import re
import fitz  # PyMuPDF
import docx
import hashlib
import numpy as np
from datetime import datetime
from collections import defaultdict
from werkzeug.utils import secure_filename
from models import db, Document, Paragraph

# Import OCR libraries if installed
try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

def save_uploaded_file(file, upload_folder):
    """Save the uploaded file to the uploads folder."""
    filename = secure_filename(file.filename)
    # Add timestamp to filename to avoid duplicates
    base, ext = os.path.splitext(filename)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_filename = f"{base}_{timestamp}{ext}"
    
    file_path = os.path.join(upload_folder, unique_filename)
    file.save(file_path)
    
    return {
        'original_filename': filename,
        'filename': unique_filename,
        'file_path': file_path,
        'file_type': ext[1:].lower()  # Remove the dot from extension
    }

def compute_paragraph_hash(text):
    """Compute a similarity hash for a paragraph to help with comparison."""
    # Normalize text: lowercase, remove punctuation, normalize whitespace
    normalized = re.sub(r'[^\w\s]', '', text.lower())
    normalized = re.sub(r'\s+', ' ', normalized).strip()
    
    # Compute hash
    return hashlib.md5(normalized.encode('utf-8')).hexdigest()

def extract_text_with_ocr(file_path):
    """Extract text from scanned PDFs using OCR."""
    if not OCR_AVAILABLE:
        raise ImportError("OCR libraries (pytesseract, pdf2image) are not installed")
    
    # Convert PDF to images
    images = convert_from_path(file_path)
    
    paragraphs = []
    paragraph_types = []
    
    for img in images:
        # Extract text using pytesseract
        text = pytesseract.image_to_string(img)
        
        # Split into paragraphs (double newlines)
        page_paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        for para in page_paragraphs:
            # Basic detection of paragraph type
            if re.match(r'^[A-Z0-9][A-Z0-9\.\s]{0,20}[A-Z]', para):  # Pattern for headings
                para_type = "heading"
            elif re.match(r'^\s*[\•\-\*\d]+[\.\)\]]\s', para):  # Pattern for list items
                para_type = "list-item"
            else:
                para_type = "regular"
            
            paragraphs.append(para)
            paragraph_types.append(para_type)
    
    return list(zip(paragraphs, paragraph_types))

def extract_text_from_pdf(file_path, use_ocr_fallback=True):
    """
    Extract text from PDF files with enhanced paragraph detection.
    
    Args:
        file_path: Path to the PDF file
        use_ocr_fallback: Whether to fall back to OCR for scanned PDFs
        
    Returns:
        List of tuples (paragraph_text, paragraph_type)
    """
    doc = fitz.open(file_path)
    
    # Check if PDF might be scanned (very little text)
    if use_ocr_fallback and OCR_AVAILABLE:
        text_amount = 0
        # Sample a few pages to check for text content
        pages_to_check = min(5, len(doc))
        for page_num in range(pages_to_check):
            page = doc.load_page(page_num)
            text_amount += len(page.get_text("text").strip())
        
        # If very little text is found, consider it a scanned document
        if text_amount < 100 * pages_to_check:
            return extract_text_with_ocr(file_path)
    
    # Process as normal PDF with text
    all_blocks = []
    
    # First pass - collect all blocks with their attributes from all pages
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        blocks = page.get_text("dict")["blocks"]
        
        for block in blocks:
            if block["type"] == 0:  # Text blocks only
                block_data = {
                    'page': page_num,
                    'block': block,
                    'bbox': block["bbox"],
                    'lines': []
                }
                
                # Extract the text and properties
                for line in block["lines"]:
                    spans = line["spans"]
                    spans_text = []
                    
                    # Line properties for classification
                    line_props = {
                        'fonts': set(),
                        'sizes': set(),
                        'font_flags': set(),  # Bold, italic, etc
                        'colors': set(),
                    }
                    
                    # Extract text and collect span properties
                    for span in spans:
                        if span["text"].strip():
                            spans_text.append(span["text"])
                            line_props['fonts'].add(span["font"])
                            line_props['sizes'].add(span["size"])
                            line_props['font_flags'].add(span.get("flags", 0))
                            line_props['colors'].add(span.get("color", 0))
                    
                    if spans_text:
                        line_text = " ".join(spans_text)
                        block_data['lines'].append({
                            'text': line_text,
                            'bbox': line["bbox"],
                            'props': line_props
                        })
                
                if block_data['lines']:
                    all_blocks.append(block_data)
    
    # Second pass - group blocks into columns based on x-coordinates
    if all_blocks:
        # Find potential column boundaries using x-coordinates of blocks
        x_coords = []
        for block in all_blocks:
            x_coords.append(block['bbox'][0])  # Left edge
            x_coords.append(block['bbox'][2])  # Right edge
        
        # Use clustering or simple thresholding to identify column boundaries
        x_coords = sorted(x_coords)
        x_diffs = [x_coords[i+1] - x_coords[i] for i in range(len(x_coords)-1)]
        
        # If we have enough data, use statistical methods to find column gaps
        if len(x_diffs) > 5:
            # Use larger gaps to identify potential column separations
            mean_diff = np.mean(x_diffs)
            std_diff = np.std(x_diffs)
            threshold = mean_diff + 1.5 * std_diff
            
            potential_column_edges = []
            for i, diff in enumerate(x_diffs):
                if diff > threshold:
                    potential_column_edges.append((x_coords[i], x_coords[i+1]))
        
        # Group blocks by column (simplified version - can be enhanced)
        # For now, we'll sort blocks by x coordinate and then by y coordinate
        all_blocks.sort(key=lambda b: (b['bbox'][0], b['bbox'][1]))
    
    # Third pass - distinguish paragraph types and build paragraphs
    paragraphs = []
    paragraph_types = []
    
    # Block grouping - group adjacent blocks that likely form a paragraph
    current_paragraph = ""
    current_para_type = ""
    current_y_bottom = None
    current_block_idx = -1
    
    for block_idx, block in enumerate(all_blocks):
        for line in block['lines']:
            line_text = line['text'].strip()
            if not line_text:
                continue
                
            # Determine line type based on properties
            is_heading = False
            is_list_item = False
            
            # Check font size and formatting
            max_size = max(line['props']['sizes']) if line['props']['sizes'] else 0
            has_bold = any(flag & 16 for flag in line['props']['font_flags'])  # Check for bold text
            
            # Check for list items or bullet points
            if re.match(r'^\s*[\•\-\*\d]+[\.\)\]]\s', line_text):
                is_list_item = True
            
            # Determine if this is likely a heading
            if has_bold or max_size > 12:
                is_heading = True
            
            # Calculate current line properties
            y_top = line['bbox'][1]
            y_bottom = line['bbox'][3]
            
            # Determine paragraph type
            if is_heading:
                line_type = "heading"
            elif is_list_item:
                line_type = "list-item"
            else:
                line_type = "regular"
            
            # Decide if we should start a new paragraph or continue the current one
            start_new_paragraph = False
            
            if current_paragraph == "":
                # First paragraph
                start_new_paragraph = False
            elif line_type != current_para_type:
                # Type changed
                start_new_paragraph = True
            elif current_block_idx != block_idx:
                # New block
                y_gap = y_top - current_y_bottom if current_y_bottom is not None else 0
                if y_gap > max_size * 0.5:  # If gap is more than half the font size
                    start_new_paragraph = True
            elif line_type == "list-item":
                # Each list item is its own paragraph
                start_new_paragraph = True
            
            # Start a new paragraph if needed
            if start_new_paragraph and current_paragraph:
                paragraphs.append(current_paragraph)
                paragraph_types.append(current_para_type)
                current_paragraph = ""
            
            # Add current line to paragraph
            if current_paragraph:
                # Check if we need a space between lines
                last_char = current_paragraph[-1]
                if last_char not in ".!?:;" and not current_paragraph.endswith("-"):
                    current_paragraph += " "
            
            current_paragraph += line_text
            current_para_type = line_type
            current_y_bottom = y_bottom
            current_block_idx = block_idx
    
    # Add the last paragraph if any
    if current_paragraph:
        paragraphs.append(current_paragraph)
        paragraph_types.append(current_para_type)
    
    # Final processing: merge short paragraphs that might have been split incorrectly
    final_paragraphs = []
    final_types = []
    
    i = 0
    while i < len(paragraphs):
        current = paragraphs[i]
        current_type = paragraph_types[i]
        
        # If this is a regular paragraph and short, and next paragraph is also regular,
        # consider merging them
        if (i < len(paragraphs) - 1 and 
            current_type == "regular" and 
            paragraph_types[i+1] == "regular" and
            len(current) < 150 and
            not current.endswith(('.', '!', '?', ':', ';'))):
            
            # Merge with next paragraph
            paragraphs[i+1] = current + " " + paragraphs[i+1]
        else:
            final_paragraphs.append(current)
            final_types.append(current_type)
        
        i += 1
    
    return list(zip(final_paragraphs, final_types))

def extract_text_from_docx(file_path):
    """Extract text from DOCX files with enhanced paragraph and structure detection."""
    doc = docx.Document(file_path)
    
    # First pass: collect all content and their properties
    elements = []
    
    # Process paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Get paragraph style and level info
        style_name = para.style.name.lower() if para.style else ""
        outline_level = para.style.style_id.startswith('Heading') if para.style else False
        
        # Check if any runs have bold formatting
        has_bold = any(run.bold for run in para.runs if hasattr(run, 'bold') and run.bold)
        
        # Check for list styles
        is_list = (
            'list' in style_name or 
            para._p.pPr is not None and 
            para._p.pPr.numPr is not None or
            text.startswith(('•', '-', '*', '○', '▪', '■')) or 
            re.match(r'^\d+[\.\)]\s', text)
        )
        
        # Determine paragraph type
        if outline_level or 'heading' in style_name or has_bold and len(text) < 200:
            para_type = "heading"
        elif is_list:
            para_type = "list-item"
        else:
            para_type = "regular"
        
        elements.append({
            'type': 'paragraph',
            'content': text,
            'para_type': para_type
        })
    
    # Process tables
    for table in doc.tables:
        # First, check if table is empty or just formatting
        table_has_content = False
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_has_content = True
                    break
            if table_has_content:
                break
        
        if not table_has_content:
            continue
        
        # Process table headers (first row) differently
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    if i == 0:  # First row - consider header
                        para_type = "table-header"
                    else:
                        para_type = "table-cell"
                    
                    elements.append({
                        'type': 'table_cell',
                        'content': f"[TABLE] {cell_text}",
                        'para_type': para_type,
                        'row': i,
                        'is_header': i == 0
                    })
    
    # Second pass: Final processing and cleanup
    result = []
    
    # Merge adjacent paragraphs of the same type if needed
    current_paragraph = ""
    current_type = ""
    
    for element in elements:
        if element['type'] == 'paragraph':
            text = element['content']
            para_type = element['para_type']
            
            # Decide whether to continue the current paragraph or start a new one
            if current_paragraph and current_type == para_type == "regular":
                # Only merge regular paragraphs that seem to continue
                if not current_paragraph.endswith(('.', '!', '?', ':', ';')):
                    current_paragraph += " " + text
                    continue
            
            # Start a new paragraph
            if current_paragraph:
                result.append((current_paragraph, current_type))
            
            current_paragraph = text
            current_type = para_type
        
        elif element['type'] == 'table_cell':
            # Always add table cells as separate elements
            if current_paragraph:
                result.append((current_paragraph, current_type))
                current_paragraph = ""
                current_type = ""
            
            result.append((element['content'], element['para_type']))
    
    # Add the last paragraph if any
    if current_paragraph:
        result.append((current_paragraph, current_type))
    
    return result

def detect_file_type(file_path):
    """Auto-detect the file type for a given file path."""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return 'pdf'
    elif ext in ['.docx', '.doc']:
        return 'docx'
    else:
        # Try to infer from content
        try:
            with open(file_path, 'rb') as f:
                header = f.read(4)
                if header == b'%PDF':
                    return 'pdf'
                # Check for DOCX (ZIP format)
                if header[:2] == b'PK':
                    return 'docx'
        except:
            pass
        
        raise ValueError(f"Unsupported or unrecognized file type: {ext}")

def extract_paragraphs(file_path):
    """
    Universal function to extract paragraphs from a document.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        List of tuples (paragraph_text, paragraph_type)
    """
    file_type = detect_file_type(file_path)
    
    if file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_type == 'docx':
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

def process_document(file_info):
    """Process document and store text in database with enhanced paragraph detection."""
    file_path = file_info['file_path']
    file_type = file_info['file_type']
    
    # Create document record
    document = Document(
        filename=file_info['filename'],
        original_filename=file_info['original_filename'],
        file_type=file_type
    )
    db.session.add(document)
    db.session.flush()  # Get ID without committing transaction
    
    # Extract text based on file type
    paragraph_data = extract_paragraphs(file_path)
    
    # Add paragraphs to database
    for idx, (text, para_type) in enumerate(paragraph_data):
        if text.strip():  # Skip empty paragraphs
            similarity_hash = compute_paragraph_hash(text)
            paragraph = Paragraph(
                document_id=document.id,
                text=text,
                paragraph_type=para_type,
                index=idx,
                similarity_hash=similarity_hash
            )
            db.session.add(paragraph)
    
    db.session.commit()
    return document.id

def find_duplicate_paragraphs():
    """Find paragraphs that are duplicated across documents."""
    from sqlalchemy import func
    
    # Query for hashes that appear in multiple documents
    duplicate_hashes = db.session.query(Paragraph.similarity_hash).filter(
        Paragraph.similarity_hash.isnot(None)
    ).group_by(Paragraph.similarity_hash).having(
        func.count(func.distinct(Paragraph.document_id)) > 1
    ).all()
    
    results = []
    for hash_value, in duplicate_hashes:
        # Get all paragraphs with this hash
        paragraphs = Paragraph.query.filter_by(similarity_hash=hash_value).all()
        
        # Group by document
        docs = defaultdict(list)
        for para in paragraphs:
            docs[para.document_id].append(para)
        
        # Only include if from multiple documents
        if len(docs) > 1:
            results.append({
                'hash': hash_value,
                'documents': docs,
                'paragraph_example': paragraphs[0].text if paragraphs else "",
                'count': len(paragraphs)
            })
    
    return results
